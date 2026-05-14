from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data" / "processed"
REPORTS = ROOT / "reports"
TEMPLATE = Path(r"C:\Users\Ctrl Diksir\Downloads\Capstone Project Report.docx")
OUTPUT = REPORTS / "final_report.docx"

DISCLAIMER = (
    "This study is an independent civic-tech and governance analytics research project "
    "based on publicly available information and evidence-based analysis. It does not "
    "represent any political party, candidate, or government institution."
)


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def style_document(doc: Document) -> None:
    set_margins(doc)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.12

    for style_name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_center(doc: Document, text: str, size: int = 12, bold: bool = False, after: int = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.25)


def cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("" if pd.isna(text) else str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], max_rows: int | None = None) -> None:
    data = df[columns].copy()
    if max_rows is not None:
        data = data.head(max_rows)

    table = doc.add_table(rows=1, cols=len(columns))
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell_shading(cell, "EDEDED")
        set_cell(cell, header, bold=True)

    for _, row in data.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.2f}".rstrip("0").rstrip(".")
            set_cell(cells[idx], value)
    doc.add_paragraph()


def strongest_categories(category_scores: pd.DataFrame) -> pd.DataFrame:
    return category_scores.sort_values(["person", "weighted_points"], ascending=[True, False]).groupby("person").head(3)


def key_evidence(evidence: pd.DataFrame) -> pd.DataFrame:
    ranked = evidence[evidence["evidence_type"].ne("insufficient_public_evidence")].copy()
    ranked["_rank"] = ranked["evidence_type"].map(
        {"quantitative": 0, "project_report": 1, "regulator_context": 2, "policy_target_context": 3, "qualitative": 4}
    ).fillna(5)
    return ranked.sort_values(["person", "_rank", "confidence"], ascending=[True, True, False]).groupby("person").head(3)


def plural(count: int, singular: str, plural_label: str | None = None) -> str:
    return f"{count} {singular if count == 1 else plural_label or singular + 's'}"


def build_report() -> Path:
    REPORTS.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE) if TEMPLATE.exists() else Document()
    clear_body(doc)
    style_document(doc)

    leader_scores = pd.read_csv(DATA / "leader_scores.csv")
    category_scores = pd.read_csv(DATA / "category_scores.csv")
    evidence = pd.read_csv(DATA / "clean_evidence_ledger.csv")
    sources = pd.read_csv(DATA / "source_verification_table.csv")

    doc.add_heading("Abstract", level=1)
    add_paragraph(
        doc,
        "This capstone report presents an evidence-based governance analytics study of selected public-sector "
        "and public-utility leadership records. The study uses publicly available source documents, regulator "
        "reports, official statements, verified media reports, and structured scoring rules to compare leadership "
        "performance across infrastructure, economy, youth empowerment, crisis handling, transparency, public "
        "utility management, and innovation. The aim is not to produce a political endorsement, but to demonstrate "
        "how data science and analytics can support transparent civic research through structured evidence, source "
        "verification, reproducible scoring, and clear reporting of limitations."
    )

    doc.add_heading("1.0 Introduction", level=1)
    add_paragraph(
        doc,
        "Governance performance is often discussed through public opinion, political affiliation, and fragmented "
        "media narratives. Such discussion can be difficult to verify because claims are rarely organized into a "
        "structured evidence model. This project applies data science and analytics techniques to civic research by "
        "building a public-impact analytics framework that compares leadership records using source-backed evidence."
    )
    add_paragraph(
        doc,
        "The research focuses on three public figures: Prof. Isa Ali Ibrahim Pantami, Dr. Jamilu Isyaku Gwamna, "
        "and Hon. Saidu Ahmed Alkali. Their records are examined through available public evidence rather than "
        "political claims. The resulting system includes a cleaned evidence ledger, source verification table, "
        "weighted scoring model, dashboard, downloadable datasets, and this final academic report."
    )

    doc.add_heading("1.1 Problem Statement", level=1)
    add_paragraph(
        doc,
        "Public leadership assessment often lacks a transparent and reproducible evidence structure. Citizens, "
        "researchers, and policy observers may find claims about governance performance, but those claims are not "
        "always linked to source quality, measurable indicators, or stated limitations. The problem addressed in "
        "this study is the absence of a simple analytical framework for comparing leadership impact using publicly "
        "available evidence while avoiding unsupported conclusions."
    )

    doc.add_heading("1.2 Aim and Objectives", level=1)
    add_paragraph(doc, "The aim of this project is to develop an evidence-based public-impact analytics framework for governance research.")
    add_bullets(
        doc,
        [
            "To collect and clean publicly available evidence about selected leadership records.",
            "To classify evidence by category, source type, confidence, and data quality.",
            "To design a scoring model that rewards verifiable and measurable evidence.",
            "To develop a dashboard for exploring leadership evidence and category scores.",
            "To produce a final academic report based on the research outputs.",
        ],
    )

    doc.add_heading("2.0 Methodology", level=1)
    add_paragraph(
        doc,
        "The study followed a structured analytics workflow consisting of data collection, data cleaning, source "
        "verification, evidence classification, score generation, dashboard implementation, and report production."
    )
    add_bullets(
        doc,
        [
            "Data Collection: public records, official publications, regulator reports, source tables, and verified media reports were gathered.",
            "Data Cleaning: names, categories, source IDs, metric values, confidence scores, and data-quality flags were standardized.",
            "Evidence Classification: each row was classified as quantitative, project report, regulator context, policy context, qualitative evidence, or insufficient public evidence.",
            "Scoring: evidence scores were aggregated by leader and category, then multiplied by category weights to calculate weighted points.",
            "Deployment: the outputs were presented in a Streamlit dashboard with downloadable datasets and reports.",
        ],
    )

    doc.add_heading("3.0 Data Collection and Description", level=1)
    add_paragraph(
        doc,
        f"The final dataset contains {len(evidence)} evidence records and {len(sources)} source records. "
        "Each evidence record includes the leader name, role period, category, indicator, metric value where available, "
        "claim summary, evidence type, confidence score, source ID, and data-quality flag."
    )
    add_table(
        doc,
        leader_scores.sort_values("overall_governance_score", ascending=False),
        ["person", "overall_governance_score", "evidence_count", "quantitative_evidence_count", "category_coverage_percent"],
        ["Leader", "Overall Score", "Evidence Count", "Quantitative Count", "Coverage %"],
    )

    doc.add_heading("4.0 Exploratory Data Analysis", level=1)
    add_paragraph(
        doc,
        "Exploratory analysis was used to understand evidence coverage, category strength, source availability, "
        "and the distribution of quantitative and qualitative records. The analysis showed that evidence coverage "
        "is uneven across leaders and categories, which is expected in public-record research."
    )
    add_table(
        doc,
        strongest_categories(category_scores),
        ["person", "category", "category_score", "weight", "weighted_points", "coverage_status"],
        ["Leader", "Category", "Score", "Weight", "Weighted Points", "Coverage Status"],
        max_rows=12,
    )

    doc.add_heading("5.0 Implementation", level=1)
    add_paragraph(
        doc,
        "The project was implemented in Python using pandas for data processing, a scoring script for evidence "
        "aggregation, and Streamlit for dashboard deployment. The dashboard provides pages for overview metrics, "
        "leadership profiles, category analysis, governance scorecards, timeline analysis, source verification, "
        "methodology, downloads, and project information."
    )
    add_paragraph(
        doc,
        "The scoring model uses category weights: Infrastructure 20%, Economy 20%, Youth Empowerment 20%, Crisis "
        "Handling 15%, Transparency 10%, Public Utility Management 10%, and Innovation 5%. These weights are used "
        "to calculate overall governance scores from category-level evidence."
    )

    doc.add_heading("5.1 Results and Analysis", level=1)
    add_paragraph(
        doc,
        "The results show differences in the type and strength of public evidence available for each leader. A higher "
        "score should be interpreted as stronger evidence coverage within the stated model, not as a complete judgment "
        "of personal character or total public value."
    )
    for _, row in leader_scores.sort_values("overall_governance_score", ascending=False).iterrows():
        add_paragraph(
            doc,
            f"{row['person']} recorded an overall evidence-backed score of {row['overall_governance_score']:.2f}/100, "
            f"with {plural(int(row['evidence_count']), 'evidence record')} and "
            f"{plural(int(row['quantitative_evidence_count']), 'quantitative evidence record')}."
        )

    doc.add_heading("5.2 Key Evidence Highlights", level=1)
    add_table(
        doc,
        key_evidence(evidence),
        ["person", "category", "indicator", "claim_summary", "source_id"],
        ["Leader", "Category", "Indicator", "Evidence Summary", "Source"],
        max_rows=12,
    )

    doc.add_heading("5.3 Limitations", level=1)
    add_bullets(
        doc,
        [
            "The research depends on publicly available evidence and does not include private or unpublished records.",
            "Sector-level outcomes may involve many institutions and should not be treated as sole-person causation.",
            "Some evidence records come from media interviews or official statements and may require further independent verification.",
            "Missing public data is treated as a research limitation, not as proof of poor performance.",
        ],
    )

    doc.add_heading("6.0 Conclusion", level=1)
    add_paragraph(
        doc,
        "This capstone project demonstrates how data science and analytics can support civic research by organizing "
        "public leadership evidence into a transparent, reproducible, and source-backed framework. The project shows "
        "that governance analytics is most useful when it clearly separates evidence, scoring logic, source quality, "
        "and research limitations. The final dashboard and report provide a foundation that can be expanded with new "
        "sources, updated metrics, and additional leaders."
    )

    doc.add_heading("References", level=1)
    for _, source in sources.iterrows():
        year = str(source.get("publication_date", "n.d."))[:4] if pd.notna(source.get("publication_date")) else "n.d."
        add_paragraph(doc, f"{source['organization']}. ({year}). {source['title']}. {source['url']}")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("Appendix", level=1)
    doc.add_heading("Appendix A: Scoring Formula", level=2)
    add_paragraph(doc, "category_score = average evidence_score for scored evidence rows")
    add_paragraph(doc, "weighted_points = category_score x category_weight")
    add_paragraph(doc, "overall_governance_score = sum of weighted_points across all categories")
    doc.add_heading("Appendix B: Research Disclaimer", level=2)
    add_paragraph(doc, DISCLAIMER)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
