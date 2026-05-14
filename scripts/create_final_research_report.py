from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data" / "processed"
REPORTS = ROOT / "reports"
OUTPUT = REPORTS / "final_public_impact_research_report.docx"

DISCLAIMER = (
    "This report is an independent civic-tech and governance analytics research output "
    "based on publicly available information and evidence-based analysis. It does not "
    "represent any political party, candidate, or government institution."
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("" if pd.isna(text) else str(text))
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], labels: list[str], max_rows: int | None = None) -> None:
    visible = df[columns].copy()
    if max_rows is not None:
        visible = visible.head(max_rows)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    for idx, label in enumerate(labels):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "EAF0F7")
        set_cell_text(cell, label, bold=True, color=RGBColor(11, 31, 58))

    for _, row in visible.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.2f}".rstrip("0").rstrip(".")
            set_cell_text(cells[idx], value)

    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for style_name, size in [("Title", 22), ("Heading 1", 15), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(11, 31, 58)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "Public Impact Analytics Research Report"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(90, 98, 112)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Final research report"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(90, 98, 112)


def person_summary(leader_scores: pd.DataFrame) -> list[str]:
    ordered = leader_scores.sort_values("overall_governance_score", ascending=False)
    rows = []
    for _, row in ordered.iterrows():
        rows.append(
            f"{row['person']} recorded an overall evidence-backed score of "
            f"{row['overall_governance_score']:.2f}/100, supported by "
            f"{int(row['evidence_count'])} evidence items and "
            f"{int(row['quantitative_evidence_count'])} measured-data items."
        )
    return rows


def top_evidence(evidence: pd.DataFrame) -> pd.DataFrame:
    ranked = evidence[evidence["evidence_type"].ne("insufficient_public_evidence")].copy()
    ranked["_rank"] = ranked["evidence_type"].map(
        {
            "quantitative": 0,
            "project_report": 1,
            "regulator_context": 2,
            "policy_target_context": 3,
            "qualitative": 4,
        }
    ).fillna(5)
    ranked = ranked.sort_values(["person", "_rank", "confidence"], ascending=[True, True, False])
    return ranked.groupby("person", as_index=False).head(3).drop(columns=["_rank"])


def build_report() -> Path:
    REPORTS.mkdir(parents=True, exist_ok=True)
    leader_scores = pd.read_csv(DATA / "leader_scores.csv")
    category_scores = pd.read_csv(DATA / "category_scores.csv")
    evidence = pd.read_csv(DATA / "clean_evidence_ledger.csv")
    sources = pd.read_csv(DATA / "source_verification_table.csv")

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Final Public Impact Analytics Research Report")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Evidence-based governance and leadership performance analysis")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(75, 85, 99)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared by Abubakar Ahmad, Data Scientist and Analytics Professional").bold = True
    doc.add_paragraph(f"Report date: {date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph(DISCLAIMER)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This final report summarizes the public-impact analytics research behind the dashboard. "
        "It compares selected leadership records using cited public evidence, category weights, "
        "source reliability, and documented limitations."
    )
    add_bullets(doc, person_summary(leader_scores))

    doc.add_heading("Overall Results", level=1)
    overall = leader_scores.sort_values("overall_governance_score", ascending=False)
    add_table(
        doc,
        overall,
        ["person", "overall_governance_score", "evidence_count", "quantitative_evidence_count", "category_coverage_percent"],
        ["Leader", "Overall Score", "Evidence Items", "Measured Data", "Coverage %"],
    )

    doc.add_heading("Category Performance", level=1)
    doc.add_paragraph(
        "The category results show where the strongest public evidence was found for each leader. "
        "A higher score means the available evidence was stronger, more measurable, or better documented."
    )
    compact_scores = category_scores.sort_values(["person", "weighted_points"], ascending=[True, False])
    add_table(
        doc,
        compact_scores,
        ["person", "category", "category_score", "weight", "weighted_points", "coverage_status"],
        ["Leader", "Category", "Score", "Weight", "Weighted Points", "Coverage"],
        max_rows=24,
    )

    doc.add_heading("Key Evidence Highlights", level=1)
    highlights = top_evidence(evidence)
    add_table(
        doc,
        highlights,
        ["person", "category", "indicator", "claim_summary", "source_id"],
        ["Leader", "Category", "Indicator", "Evidence Summary", "Source"],
        max_rows=12,
    )

    doc.add_heading("Methodology Summary", level=1)
    add_bullets(
        doc,
        [
            "Evidence was grouped by leader and governance category.",
            "Measured quantitative data received the strongest score because it can be verified more directly.",
            "Project reports, regulator context, and policy context received moderate scores.",
            "Descriptive or interview-based evidence received lower scores unless supported by stronger documentation.",
            "Rows with insufficient public evidence were kept as research gaps and did not add positive points.",
        ],
    )

    doc.add_heading("Research Limitations", level=1)
    add_bullets(
        doc,
        [
            "Public records are uneven across leaders and sectors.",
            "Sector-level outcomes should not be treated as the sole achievement of one person.",
            "Some evidence comes from media reports or official statements that may require further verification.",
            "Missing public data is treated as a limitation, not as proof of poor performance.",
        ],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("Source Register", level=1)
    add_table(
        doc,
        sources,
        ["source_id", "title", "organization", "source_type", "reliability_tier", "url"],
        ["ID", "Title", "Organization", "Type", "Reliability", "URL"],
        max_rows=30,
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
