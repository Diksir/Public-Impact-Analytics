from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data" / "processed"
REPORTS = ROOT / "reports"
DISCLAIMER = (
    "This platform is an independent civic-tech and governance analytics research "
    "project based on publicly available information and evidence-based analysis. "
    "It does not represent any political party, candidate, or government institution."
)


def style_document(doc: Document, title: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 22, RGBColor(11, 31, 58)),
        ("Heading 1", 15, RGBColor(11, 31, 58)),
        ("Heading 2", 12, RGBColor(11, 31, 58)),
    ]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = color

    p = doc.add_paragraph()
    p.style = "Title"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title).bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 85, 99)
    doc.add_paragraph(DISCLAIMER)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], max_rows: int = 12) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(columns):
        hdr[idx].text = col
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(11, 31, 58)
    for _, row in df[columns].head(max_rows).iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = "" if pd.isna(row[col]) else str(row[col])
            cells[idx].text = value[:900]


def apa(source) -> str:
    year = str(source["publication_date"])[:4] if pd.notna(source["publication_date"]) else "n.d."
    org = source["organization"]
    title = source["title"]
    url = source["url"]
    return f"{org}. ({year}). {title}. {url}"


def mla(source) -> str:
    org = source["organization"]
    title = source["title"]
    date = source["publication_date"] if pd.notna(source["publication_date"]) else "n.d."
    url = source["url"]
    return f'{org}. "{title}." {date}, {url}.'


def governance_report() -> None:
    leader_scores = pd.read_csv(DATA / "leader_scores.csv")
    category_scores = pd.read_csv(DATA / "category_scores.csv")
    evidence = pd.read_csv(DATA / "clean_evidence_ledger.csv")
    timeline = pd.read_csv(DATA / "timeline.csv")

    doc = Document()
    style_document(
        doc,
        "Governance Performance Analytics Report",
        "Evidence-based comparison of selected public-sector and public-utility leadership indicators",
    )
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report compares documented governance and institutional performance indicators. "
        "The scores are derived from cited evidence rows and should be read as a civic-tech analytic aid, not a definitive judgment."
    )
    add_table(
        doc,
        leader_scores,
        [
            "person",
            "overall_governance_score",
            "evidence_count",
            "quantitative_evidence_count",
            "category_coverage_percent",
        ],
    )

    doc.add_heading("Category Scorecards", level=1)
    add_table(
        doc,
        category_scores,
        ["person", "category", "weight", "category_score", "evidence_count", "coverage_status"],
        max_rows=20,
    )

    doc.add_heading("Key Evidence Ledger", level=1)
    add_table(
        doc,
        evidence,
        ["person", "category", "indicator", "metric_value", "metric_unit", "claim_summary", "source_id"],
        max_rows=18,
    )

    doc.add_heading("Leadership Journey Timeline", level=1)
    add_table(
        doc,
        timeline,
        ["candidate_name", "life_stage", "year_or_period", "role_or_event", "institution", "citation", "impact_score"],
        max_rows=25,
    )

    doc.add_heading("Limitations", level=1)
    for item in [
        "Sector-level outcomes are not treated as sole-person causation.",
        "Media interviews are useful public records but carry lower confidence than regulator datasets.",
        "Missing public data is reported as a limitation and is not interpreted as poor performance.",
        "Users should verify source links and update datasets before formal publication.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.save(REPORTS / "governance_report.docx")


def methodology_report() -> None:
    weights = pd.read_csv(DATA / "scoring_weights.csv")
    doc = Document()
    style_document(doc, "Governance Analytics Methodology", "Data cleaning, KPI generation, scoring formula, and interpretation guidance")
    doc.add_heading("Data Model", level=1)
    doc.add_paragraph(
        "The platform uses two linked research datasets: an evidence ledger for scored governance indicators and a leadership journey timeline for full life-cycle candidate profiles."
    )
    doc.add_paragraph(
        "The leadership journey schema is: candidate_name, life_stage, year_or_period, role_or_event, institution, sector, achievement, evidence, citation, impact_category, and impact_score."
    )
    doc.add_heading("Cleaning Rules", level=1)
    for item in [
        "Standardize leader names and category labels.",
        "Preserve missing metric values as blanks rather than replacing them with fabricated values.",
        "Join evidence rows to source records by source_id.",
        "Flag quantitative, contextual, and insufficient-public-evidence rows separately.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Scoring Formula", level=1)
    doc.add_paragraph("Evidence score is based on evidence type and confidence:")
    doc.add_paragraph("Quantitative rows receive the strongest base score; project reports and regulator context receive medium base scores; qualitative interview or narrative rows receive lower base scores.")
    doc.add_paragraph("category_score = average(evidence_score for scored rows in that category)")
    doc.add_paragraph("weighted_points = category_score * category_weight")
    doc.add_paragraph("overall_governance_score = sum(weighted_points across all categories)")
    add_table(doc, weights, ["category", "weight", "definition"], max_rows=10)
    doc.add_heading("Interpretation", level=1)
    doc.add_paragraph(
        "The model rewards verifiable, source-backed evidence and penalizes missing coverage by not awarding positive points for unsupported categories. It is not a popularity model and does not measure private virtue."
    )
    doc.save(REPORTS / "methodology.docx")


def references_report() -> None:
    sources = pd.read_csv(DATA / "source_verification_table.csv")
    doc = Document()
    style_document(doc, "References And Source Verification", "APA citations, MLA citations, bibliography, and source audit table")
    doc.add_heading("APA References", level=1)
    for _, source in sources.iterrows():
        doc.add_paragraph(apa(source), style="List Number")
    doc.add_heading("MLA References", level=1)
    for _, source in sources.iterrows():
        doc.add_paragraph(mla(source), style="List Number")
    doc.add_heading("Source Verification Table", level=1)
    add_table(
        doc,
        sources,
        ["source_id", "title", "organization", "source_type", "publication_date", "reliability_tier", "url"],
        max_rows=30,
    )
    doc.save(REPORTS / "references.docx")


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    governance_report()
    methodology_report()
    references_report()


if __name__ == "__main__":
    main()
